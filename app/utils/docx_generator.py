"""DOCX generation using python-docx with atomic writes and markdown parsing."""

from __future__ import annotations

import re
from typing import Any

import structlog
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement

from app.utils.file_manager import cleanup_file, finalize_temp_file, get_temp_filepath

log = structlog.get_logger(__name__)


def _set_document_properties(doc: Document, metadata: dict[str, Any]) -> None:
    core = doc.core_properties
    core.title = metadata.get("title", "")
    core.subject = metadata.get("subject", "")
    core.keywords = metadata.get("keywords", "")


def _set_page_layout(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4
    section.page_width = Inches(8.27)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)


def _apply_run_formatting(run: Any, text: str) -> None:
    """Parse inline bold/italic markers and set run properties."""
    # Strip outer markers already handled by caller
    run.text = text


def _add_formatted_paragraph(doc: Document, line: str, style: str = "Normal") -> None:
    """Add a paragraph with inline bold/italic parsing."""
    try:
        para = doc.add_paragraph(style=style)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para.paragraph_format.line_spacing = 1.15

        # Tokenise **bold**, *italic*
        pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*)")
        parts = pattern.split(line)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.bold = True
                run.font.name = "Calibri"
                run.font.size = Pt(11)
            elif part.startswith("*") and part.endswith("*"):
                run = para.add_run(part[1:-1])
                run.italic = True
                run.font.name = "Calibri"
                run.font.size = Pt(11)
            else:
                run = para.add_run(part)
                run.font.name = "Calibri"
                run.font.size = Pt(11)
    except Exception as exc:
        log.warning("docx_paragraph_error", line=line[:80], error=str(exc))


def generate_docx(
    content: str,
    filepath: str,
    metadata: dict[str, Any],
) -> None:
    """Generate a DOCX file from markdown-style content with atomic write."""
    temp_path = get_temp_filepath(filepath)
    try:
        doc = Document()
        _set_page_layout(doc)
        _set_document_properties(doc, metadata)

        for line in content.splitlines():
            try:
                stripped = line.rstrip()
                if not stripped:
                    doc.add_paragraph()
                    continue

                if stripped.startswith("#### "):
                    p = doc.add_heading(stripped[5:], level=4)
                elif stripped.startswith("### "):
                    p = doc.add_heading(stripped[4:], level=3)
                elif stripped.startswith("## "):
                    p = doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("# "):
                    p = doc.add_heading(stripped[2:], level=1)
                elif stripped.startswith("- ") or stripped.startswith("* "):
                    _add_formatted_paragraph(doc, stripped[2:], style="List Bullet")
                elif re.match(r"^\d+\.\s", stripped):
                    text = re.sub(r"^\d+\.\s", "", stripped)
                    _add_formatted_paragraph(doc, text, style="List Number")
                else:
                    _add_formatted_paragraph(doc, stripped)
            except Exception as exc:
                log.warning("docx_line_error", error=str(exc))

        doc.save(temp_path)
        finalize_temp_file(temp_path, filepath)
        log.info("docx_generated", filepath=filepath)
    except Exception as exc:
        cleanup_file(temp_path)
        raise exc
